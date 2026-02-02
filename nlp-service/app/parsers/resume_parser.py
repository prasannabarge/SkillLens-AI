"""
Resume Parser
Extracts text from various resume file formats
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class ResumeParser:
    """Parses resume files and extracts text content"""
    
    def parse(self, file_bytes: bytes) -> str:
        """
        Parse resume file and extract text
        
        Args:
            file_bytes: Raw file bytes
            
        Returns:
            Extracted text content
        """
        # Try different parsers based on file signature
        text = None
        
        # Try PDF first
        if file_bytes[:4] == b'%PDF':
            text = self._parse_pdf(file_bytes)
        
        # Try DOCX
        elif file_bytes[:4] == b'PK\x03\x04':
            text = self._parse_docx(file_bytes)
        
        # Try plain text
        else:
            text = self._parse_text(file_bytes)
        
        if text:
            # Clean up text
            text = self._clean_text(text)
            
        return text or ""
    
    def _parse_pdf(self, file_bytes: bytes) -> Optional[str]:
        """Parse PDF file"""
        try:
            import pdfplumber
            
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return "\n".join(text_parts)
        except ImportError:
            logger.warning("pdfplumber not installed, trying PyPDF2")
            return self._parse_pdf_fallback(file_bytes)
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return self._parse_pdf_fallback(file_bytes)
    
    def _parse_pdf_fallback(self, file_bytes: bytes) -> Optional[str]:
        """Fallback PDF parser using PyPDF2"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF fallback parsing error: {e}")
            return None
    
    def _parse_docx(self, file_bytes: bytes) -> Optional[str]:
        """Parse DOCX file"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_bytes))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return None
    
    def _parse_text(self, file_bytes: bytes) -> Optional[str]:
        """Parse plain text file"""
        try:
            # Try UTF-8 first
            try:
                return file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                # Try other encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        return file_bytes.decode(encoding)
                    except UnicodeDecodeError:
                        continue
            return None
        except Exception as e:
            logger.error(f"Text parsing error: {e}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        import re
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\-\+\#\@\(\)]', ' ', text)
        
        # Normalize whitespace again
        text = ' '.join(text.split())
        
        return text.strip()
